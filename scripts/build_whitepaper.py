from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/livo-architecture-whitepaper.docx")
PDF_OUT = Path("docs/livo-architecture-whitepaper.pdf")

CONTENT = {
    "title": "Personal Lang Coach",
    "subtitle": "System Architecture White Paper | Livo AI SWE Assessment | July 2026",
    "overview": (
        "Personal Lang Coach is a Vercel-hosted static web app for 15-45 second English speech uploads. "
        "The shipped design intentionally keeps audio processing inside the browser: Vercel serves HTML, CSS, and JavaScript, while the Web Audio API decodes the recording and the local analyzer produces scoring, mistake ranges, and coaching notes."
    ),
    "diagram": [
        "User browser",
        "  -> file duration gate + explicit consent",
        "  -> Web Audio decoder",
        "  -> analyzer.js: voicing, pause, clarity, rhythm, pace",
        "  -> results renderer: score, waveform, segment/word highlights",
        "Vercel static hosting serves assets only; no audio API or object store is used by default.",
    ],
    "models": (
        "The implementation uses the browser Web Audio API and deterministic acoustic heuristics, not an external STT, TTS, LLM, or phoneme API. This was chosen over Whisper, Deepgram, or cloud LLM judging because the assessment must be reachable without private API keys and because local processing materially reduces DPDP risk for voice data."
    ),
    "scoring": (
        "Audio is split into 30 ms frames with 15 ms hops. RMS energy estimates voicing and clarity; zero-crossing rate flags noisy consonant-like spans; pause gaps flag rhythm breaks. The overall score weights clarity 48%, rhythm 27%, pace 20%, and duration compliance 5%."
    ),
    "highlighting": (
        "The app returns time-coded issues such as unclear segment, noisy articulation, uneven stress, long pause, or low speech density. If the learner supplies the expected transcript, issue time ranges are mapped onto words for word-level highlights; otherwise the same findings appear as segment-level coaching cards."
    ),
    "dpdp": (
        "Voice recordings can identify a speaker and are treated as digital personal data. The app therefore uses data minimization by default: it does not upload, store, or log the audio. Processing happens only after an affirmative consent checkbox, and reset/page refresh clears the local object URL, decoded buffer, transcript, and analysis state."
    ),
    "storage": (
        "No server-side audio storage exists in the default deployment. Retention is limited to browser memory for the active session; the specified purpose ends when feedback is shown or the user resets."
    ),
    "residency": (
        "Because audio is not transmitted, there is no audio data residency dependency. Vercel only hosts static assets. If a production STT service is added, it should use a contracted processor, approved transfer/residency posture, encryption, short TTL storage, deletion APIs, and logs that exclude raw audio."
    ),
    "rights": (
        "The UI exposes plain-language notice, purpose limitation, and deletion by reset. A production account-based version would add an authenticated deletion request flow, grievance contact, and audit trail."
    ),
    "tradeoffs": (
        "The major trade-off is privacy and deployability over phoneme-level accuracy: the current app cannot prove a specific word was mispronounced without a transcript or STT alignment. It still provides useful learner feedback on fluency, pauses, clarity, pacing, and unstable stress. With another week, I would add optional server-side forced alignment, phoneme comparison against a target passage, calibrated CEFR-style scoring, signed transient uploads, and a teacher review dashboard."
    ),
    "source": (
        "Compliance reference: Digital Personal Data Protection Act, 2023, Gazette of India, Act No. 22 of 2023. Relevant controls: notice and consent, security safeguards, erasure, and cross-border transfer restrictions."
    ),
}


def set_font(run, name="Calibri", size=11, bold=False, color="17201F"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, before=0, after=6, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=12 if level == 1 else 8, after=5)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color="23595B")
    return paragraph


def add_body(doc, text, after=5):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_compact_bullet(doc, label, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph(paragraph, after=3, line=1.08)
    run = paragraph.add_run(f"{label}: ")
    set_font(run, bold=True)
    run = paragraph.add_run(text)
    set_font(run)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=2, after=8, line=1.0)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", size=9.5, color="1F3A3D")
        if index < len(lines) - 1:
            run.add_break()
    shade(paragraph, "EEF4F3")
    border(paragraph, "C7D8D6")


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def border(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "3")
        element.set(qn("w:color"), color)
        borders.append(element)
    p_pr.append(borders)


def set_document_styles(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    set_paragraph(title, after=2, line=1.0)
    run = title.add_run(CONTENT["title"])
    set_font(run, size=22, bold=True, color="17201F")

    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, after=10)
    run = subtitle.add_run(CONTENT["subtitle"])
    set_font(run, size=10.5, color="66736F")

    add_heading(doc, "1. Product And Architecture", 1)
    add_body(doc, CONTENT["overview"])
    add_code_block(doc, CONTENT["diagram"])

    add_heading(doc, "2. Models, APIs, Scoring, And Highlights", 1)
    add_compact_bullet(
        doc,
        "Models and APIs",
        CONTENT["models"]
    )
    add_compact_bullet(
        doc,
        "Scoring",
        CONTENT["scoring"]
    )
    add_compact_bullet(
        doc,
        "Highlighting",
        CONTENT["highlighting"]
    )

    add_heading(doc, "3. DPDP Compliance Posture", 1)
    add_body(doc, CONTENT["dpdp"])
    add_compact_bullet(
        doc,
        "Storage and retention",
        CONTENT["storage"]
    )
    add_compact_bullet(
        doc,
        "Residency and processors",
        CONTENT["residency"]
    )
    add_compact_bullet(
        doc,
        "Rights and deletion",
        CONTENT["rights"]
    )

    add_heading(doc, "4. Trade-offs And Next Week", 1)
    add_body(doc, CONTENT["tradeoffs"])

    source = doc.add_paragraph()
    set_paragraph(source, before=4, after=0, line=1.0)
    run = source.add_run(CONTENT["source"])
    set_font(run, size=8.5, color="66736F")

    doc.core_properties.title = "Personal Lang Coach Architecture White Paper"
    doc.core_properties.subject = "Livo AI SWE Assessment"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUT)
    build_pdf()


def build_pdf():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=12,
        spaceAfter=5,
        textColor=colors.HexColor("#17201F"),
    )
    title = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=25,
        alignment=0,
        spaceAfter=2,
        textColor=colors.HexColor("#17201F"),
    )
    subtitle = ParagraphStyle(
        "Subtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        spaceAfter=8,
        textColor=colors.HexColor("#66736F"),
    )
    heading = ParagraphStyle(
        "Heading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=15,
        spaceBefore=7,
        spaceAfter=4,
        textColor=colors.HexColor("#23595B"),
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=12,
        firstLineIndent=-12,
        spaceAfter=3,
    )
    code = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8.6,
        leading=10.5,
        borderColor=colors.HexColor("#C7D8D6"),
        borderWidth=0.5,
        borderPadding=6,
        backColor=colors.HexColor("#EEF4F3"),
        textColor=colors.HexColor("#1F3A3D"),
        spaceAfter=7,
    )
    source = ParagraphStyle(
        "Source",
        parent=body,
        fontSize=7.8,
        leading=9.2,
        textColor=colors.HexColor("#66736F"),
    )

    story = [
        Paragraph(CONTENT["title"], title),
        Paragraph(CONTENT["subtitle"], subtitle),
        Paragraph("1. Product And Architecture", heading),
        Paragraph(CONTENT["overview"], body),
        Preformatted("\n".join(CONTENT["diagram"]), code),
        Paragraph("2. Models, APIs, Scoring, And Highlights", heading),
        Paragraph(f"<b>Models and APIs:</b> {CONTENT['models']}", bullet),
        Paragraph(f"<b>Scoring:</b> {CONTENT['scoring']}", bullet),
        Paragraph(f"<b>Highlighting:</b> {CONTENT['highlighting']}", bullet),
        Paragraph("3. DPDP Compliance Posture", heading),
        Paragraph(CONTENT["dpdp"], body),
        Paragraph(f"<b>Storage and retention:</b> {CONTENT['storage']}", bullet),
        Paragraph(f"<b>Residency and processors:</b> {CONTENT['residency']}", bullet),
        Paragraph(f"<b>Rights and deletion:</b> {CONTENT['rights']}", bullet),
        Paragraph("4. Trade-offs And Next Week", heading),
        Paragraph(CONTENT["tradeoffs"], body),
        Spacer(1, 3),
        Paragraph(CONTENT["source"], source),
    ]
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title="Personal Lang Coach Architecture White Paper",
        author="",
    )
    doc.build(story)


if __name__ == "__main__":
    build()
